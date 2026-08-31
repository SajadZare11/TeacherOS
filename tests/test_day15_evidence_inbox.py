from __future__ import annotations

import io
import json
import os
import sys
import tempfile
import unittest
from datetime import datetime, timedelta, timezone
from pathlib import Path
from types import SimpleNamespace

PROJECT_ROOT = Path(__file__).resolve().parents[1]
BACKEND_DIR = PROJECT_ROOT / "backend"
sys.path.insert(0, str(BACKEND_DIR))
os.environ.setdefault("TELEGRAM_BOT_TOKEN", "offline-day15-token")
os.environ.setdefault("OPENROUTER_API_KEY", "offline-day15-key")

from docx import Document

import database
from class_service import create_class
from evidence_keyboards import (
    evidence_batch_details_keyboard,
    evidence_delete_confirm_keyboard,
    evidence_inbox_keyboard,
    evidence_item_view_keyboard,
    evidence_retention_keyboard,
    evidence_submission_method_keyboard,
    evidence_type_keyboard,
)
from evidence_service import (
    ClassEvidenceDisabledError,
    delete_evidence_batch,
    delete_evidence_item,
    get_evidence_batch,
    list_evidence_batches,
    parse_docx_bytes,
    parse_txt_bytes,
    purge_expired_evidence,
    sanitize_filename,
    split_evidence_text,
    submit_evidence_batch,
    update_evidence_item_label,
    validate_file_submission,
)
from feature_flags import FEATURE_ENV_VARS


def _teacher(identifier: int, username: str = "teacher") -> SimpleNamespace:
    return SimpleNamespace(
        id=identifier,
        username=f"{username}_{identifier}",
        first_name="Evidence",
        last_name="Teacher",
        language_code="en",
    )


class Day15EvidenceInboxTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.db_path = Path(self.temp_dir.name) / "test_teacheros.db"
        self.original_db = database.DATABASE_PATH
        database.DATABASE_PATH = self.db_path

        self.previous_flags = {
            name: os.environ.get(name) for name in FEATURE_ENV_VARS.values()
        }
        for name in FEATURE_ENV_VARS.values():
            os.environ[name] = "false"
        os.environ[FEATURE_ENV_VARS["classes"]] = "true"
        os.environ[FEATURE_ENV_VARS["continuity"]] = "true"
        os.environ[FEATURE_ENV_VARS["evidence"]] = "true"

        database.initialize_database(self.db_path)
        self.teacher_a = _teacher(150_001, "teacher_a")
        self.teacher_b = _teacher(150_002, "teacher_b")

        self.class_a = create_class(
            telegram_user=self.teacher_a,
            display_name="IELTS Writing Prep",
            level="B2",
            age_group="adults",
            learner_count_band="13_20",
            goal="Essay coherence and task response",
            database_path=self.db_path,
        )

    def tearDown(self) -> None:
        database.DATABASE_PATH = self.original_db
        for name, val in self.previous_flags.items():
            if val is None:
                os.environ.pop(name, None)
            else:
                os.environ[name] = val
        self.temp_dir.cleanup()

    def test_schema_v15_is_idempotent_and_creates_tables(self) -> None:
        with database.database_connection(self.db_path) as conn:
            batches_info = conn.execute("PRAGMA table_info(evidence_batches)").fetchall()
            items_info = conn.execute("PRAGMA table_info(evidence_items)").fetchall()
            batch_cols = {row["name"] for row in batches_info}
            item_cols = {row["name"] for row in items_info}

            self.assertIn("batch_uuid", batch_cols)
            self.assertIn("evidence_type", batch_cols)
            self.assertIn("retention_policy", batch_cols)
            self.assertIn("privacy_confirmed", batch_cols)

            self.assertIn("student_label", item_cols)
            self.assertIn("char_count", item_cols)
            self.assertIn("word_count", item_cols)
            self.assertIn("status", item_cols)

            version_rows = conn.execute(
                "SELECT version FROM schema_versions WHERE version = 15"
            ).fetchall()
            self.assertEqual(len(version_rows), 1)

    def test_pasted_text_splits_multi_student_responses_with_anonymous_labels(self) -> None:
        sample_text = (
            "Student 1: In my opinion, renewable energy is essential for reducing carbon emissions.\n"
            "Student 2: Although electric cars are popular, their batteries require mining rare metals.\n"
            "Student 3: Governments should invest more in public transit rather than highways.\n"
        )
        batch = submit_evidence_batch(
            telegram_user=self.teacher_a,
            class_id=self.class_a["id"],
            evidence_type="writing",
            raw_text=sample_text,
            retention_policy="30_days",
            privacy_confirmed=True,
            database_path=self.db_path,
        )

        self.assertEqual(batch["active_item_count"], 3)
        self.assertEqual(batch["evidence_type"], "writing")
        labels = [item["student_label"] for item in batch["items"]]
        self.assertEqual(labels, ["Student 1", "Student 2", "Student 3"])
        for item in batch["items"]:
            self.assertGreater(item["char_count"], 20)
            self.assertGreater(item["word_count"], 4)

    def test_separator_lines_split_cleanly(self) -> None:
        text = (
            "First essay on climate change policies and economic impacts.\n"
            "---\n"
            "Second essay examining deforestation in tropical rainforest regions.\n"
            "===\n"
            "Third essay discussing urban planning and sustainable architecture.\n"
        )
        items = split_evidence_text(text)
        self.assertEqual(len(items), 3)
        self.assertEqual(items[0]["student_label"], "Student 1")
        self.assertEqual(items[1]["student_label"], "Student 2")
        self.assertEqual(items[2]["student_label"], "Student 3")

    def test_txt_file_upload_parsing_and_binary_rejection(self) -> None:
        txt_bytes = "Student A: Travel broadens the mind.\nStudent B: Learning languages builds empathy.".encode("utf-8")
        fmt, fname, items = validate_file_submission("homework_task.txt", txt_bytes)
        self.assertEqual(fmt, "txt_file")
        self.assertEqual(fname, "homework_task.txt")
        self.assertEqual(len(items), 2)
        self.assertEqual(items[0]["student_label"], "Student A")
        self.assertEqual(items[1]["student_label"], "Student B")

        # Binary null bytes rejection
        binary_bytes = b"Student 1: Hello" + bytes([0, 1, 2]) + b"Corrupt Data"
        with self.assertRaises(ValueError):
            validate_file_submission("bad.txt", binary_bytes)

    def test_docx_file_upload_parsing_and_corrupt_file_handling(self) -> None:
        doc = Document()
        doc.add_heading("Class Homework Submissions", level=1)
        doc.add_paragraph("Student 1: Artificial intelligence will reshape healthcare diagnostics.")
        doc.add_paragraph("Student 2: Telemedicine allows rural patients to access specialist care.")

        stream = io.BytesIO()
        doc.save(stream)
        docx_bytes = stream.getvalue()

        batch = submit_evidence_batch(
            telegram_user=self.teacher_a,
            class_id=self.class_a["id"],
            evidence_type="homework_task",
            file_name="submissions.docx",
            file_bytes=docx_bytes,
            retention_policy="7_days",
            privacy_confirmed=True,
            database_path=self.db_path,
        )

        self.assertEqual(batch["source_format"], "docx_file")
        self.assertEqual(batch["source_filename"], "submissions.docx")
        self.assertEqual(batch["active_item_count"], 2)

        # Corrupt docx bytes
        with self.assertRaises(ValueError):
            parse_docx_bytes(b"PK\x03\x04not-a-valid-zip-docx")

    def test_deferred_formats_fail_safely_with_explanation(self) -> None:
        for ext in (".pdf", ".mp3", ".jpg", ".png"):
            with self.assertRaises(ValueError) as ctx:
                validate_file_submission(f"sample{ext}", b"data")
            self.assertIn("deferred", str(ctx.exception).lower())

    def test_file_size_and_batch_limits(self) -> None:
        # Oversized file (>2 MB)
        oversized = b"a" * (2 * 1024 * 1024 + 10)
        with self.assertRaises(ValueError) as ctx:
            parse_txt_bytes(oversized)
        self.assertIn("2 MB limit", str(ctx.exception))

        # Empty file
        with self.assertRaises(ValueError):
            parse_txt_bytes(b"")

        # Empty content string
        with self.assertRaises(ValueError):
            split_evidence_text("   ")

    def test_multi_tenant_isolation_and_cross_access_denial(self) -> None:
        batch = submit_evidence_batch(
            telegram_user=self.teacher_a,
            class_id=self.class_a["id"],
            evidence_type="writing",
            raw_text="Student 1: Teacher A essay.",
            database_path=self.db_path,
        )

        # Teacher B attempts to view Teacher A's batch
        b_view = get_evidence_batch(
            telegram_user_id=self.teacher_b.id,
            batch_id=batch["id"],
            database_path=self.db_path,
        )
        self.assertIsNone(b_view)

        # Teacher B attempts to list batches for Teacher A's class
        b_list = list_evidence_batches(
            telegram_user_id=self.teacher_b.id,
            class_id=self.class_a["id"],
            database_path=self.db_path,
        )
        self.assertEqual(b_list, [])

        # Teacher B attempts to update or delete Teacher A's item
        item_id = batch["items"][0]["id"]
        update_res = update_evidence_item_label(
            telegram_user_id=self.teacher_b.id,
            item_id=item_id,
            new_label="Hacked Label",
            database_path=self.db_path,
        )
        self.assertIsNone(update_res)

        del_res = delete_evidence_item(
            telegram_user_id=self.teacher_b.id,
            item_id=item_id,
            database_path=self.db_path,
        )
        self.assertFalse(del_res)

        # Confirm Teacher A's item remains intact
        a_view = get_evidence_batch(
            telegram_user_id=self.teacher_a.id,
            batch_id=batch["id"],
            database_path=self.db_path,
        )
        self.assertEqual(a_view["items"][0]["student_label"], "Student 1")

    def test_update_student_label_and_privacy_confirmation(self) -> None:
        batch = submit_evidence_batch(
            telegram_user=self.teacher_a,
            class_id=self.class_a["id"],
            evidence_type="speaking_notes",
            raw_text="Student 1: Fluent interaction during pair work simulation.",
            database_path=self.db_path,
        )
        item_id = batch["items"][0]["id"]

        updated = update_evidence_item_label(
            telegram_user=self.teacher_a.id,
            item_id=item_id,
            new_label="Pair A (Simulation)",
            database_path=self.db_path,
        )
        self.assertIsNotNone(updated)
        self.assertEqual(updated["student_label"], "Pair A (Simulation)")

        # Unconfirmed privacy submission is rejected
        with self.assertRaises(ValueError):
            submit_evidence_batch(
                telegram_user=self.teacher_a,
                class_id=self.class_a["id"],
                evidence_type="writing",
                raw_text="Student 1: Test",
                privacy_confirmed=False,
                database_path=self.db_path,
            )

    def test_item_deletion_and_batch_deletion_cascade(self) -> None:
        batch = submit_evidence_batch(
            telegram_user=self.teacher_a,
            class_id=self.class_a["id"],
            evidence_type="quiz_exit_ticket",
            raw_text="Student 1: Answer A\n---\nStudent 2: Answer B",
            database_path=self.db_path,
        )
        batch_id = batch["id"]
        item1_id = batch["items"][0]["id"]

        # Delete single item
        ok = delete_evidence_item(
            telegram_user_id=self.teacher_a.id,
            item_id=item1_id,
            database_path=self.db_path,
        )
        self.assertTrue(ok)

        batch_refreshed = get_evidence_batch(
            telegram_user_id=self.teacher_a.id,
            batch_id=batch_id,
            database_path=self.db_path,
        )
        self.assertEqual(batch_refreshed["active_item_count"], 1)

        # Delete entire batch
        b_ok = delete_evidence_batch(
            telegram_user_id=self.teacher_a.id,
            batch_id=batch_id,
            database_path=self.db_path,
        )
        self.assertTrue(b_ok)

        batches = list_evidence_batches(
            telegram_user_id=self.teacher_a.id,
            class_id=self.class_a["id"],
            database_path=self.db_path,
        )
        self.assertEqual(batches, [])

    def test_retention_policy_purge_expired(self) -> None:
        batch = submit_evidence_batch(
            telegram_user=self.teacher_a,
            class_id=self.class_a["id"],
            evidence_type="writing",
            raw_text="Student 1: Expired evidence content.",
            retention_policy="7_days",
            database_path=self.db_path,
        )
        batch_id = batch["id"]

        # Backdate the batch creation timestamp to 10 days ago
        ten_days_ago = (datetime.now(timezone.utc) - timedelta(days=10)).strftime("%Y-%m-%dT%H:%M:%S.%fZ")
        with database.database_connection(self.db_path) as conn:
            conn.execute(
                "UPDATE evidence_batches SET created_at = ? WHERE id = ?",
                (ten_days_ago, batch_id),
            )

        purged_count = purge_expired_evidence(database_path=self.db_path)
        self.assertGreaterEqual(purged_count, 1)

        with database.database_connection(self.db_path) as conn:
            item = conn.execute(
                "SELECT content, status FROM evidence_items WHERE batch_id = ?", (batch_id,)
            ).fetchone()
            self.assertEqual(item["status"], "purged")
            self.assertEqual(item["content"], "[PURGED_BY_RETENTION_POLICY]")

    def test_zero_raw_evidence_in_logs_and_telemetry(self) -> None:
        sensitive_snippet = "UniqueSecretConfidentialStudentWriting12345"
        batch = submit_evidence_batch(
            telegram_user=self.teacher_a,
            class_id=self.class_a["id"],
            evidence_type="writing",
            raw_text=f"Student 1: {sensitive_snippet}",
            database_path=self.db_path,
        )

        with database.database_connection(self.db_path) as conn:
            events = conn.execute(
                "SELECT event_name, properties_json FROM product_events WHERE class_id = ?",
                (self.class_a["id"],),
            ).fetchall()

            for event in events:
                props = str(event["properties_json"])
                self.assertNotIn(sensitive_snippet, props)
                parsed = json.loads(props)
                self.assertIn("item_count", parsed)
                self.assertIn("evidence_type", parsed)

    def test_keyboards_are_compact_and_within_64_bytes(self) -> None:
        sample_batch = [{"id": 101, "active_items": 4, "evidence_type": "writing", "created_at": "2026-09-01"}]
        sample_items = [{"id": 201, "student_label": "Student 1", "word_count": 50}]

        kbs = [
            evidence_inbox_keyboard(1234, 5, sample_batch),
            evidence_type_keyboard(1234, 5),
            evidence_retention_keyboard(1234, "w", 5),
            evidence_submission_method_keyboard(1234, "w", "30", 5),
            evidence_batch_details_keyboard(101, 1234, 5, sample_items),
            evidence_item_view_keyboard(201, 101, 5),
            evidence_delete_confirm_keyboard(101, 1234, 5),
        ]

        for kb in kbs:
            for row in kb.inline_keyboard:
                for btn in row:
                    self.assertLessEqual(
                        len(btn.callback_data.encode("utf-8")),
                        64,
                        f"Button callback data exceeds 64 bytes: {btn.callback_data}",
                    )


if __name__ == "__main__":
    unittest.main()
