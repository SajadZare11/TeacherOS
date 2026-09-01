from __future__ import annotations

import io
import json
import logging
import os
import re
import secrets
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from database import database_connection, ensure_database_user
from feature_flags import feature_enabled


logger = logging.getLogger(__name__)

MAX_FILE_BYTES = 2 * 1024 * 1024  # 2 MB
MAX_BATCH_ITEMS = 50
MAX_ITEM_CHARS = 25_000
MIN_ITEM_CHARS = 2

EVIDENCE_TYPES = (
    "writing",
    "speaking_notes",
    "quiz_exit_ticket",
    "homework_task",
    "general_work",
)
EVIDENCE_TYPE_LABELS = {
    "writing": "Writing Sample",
    "speaking_notes": "Speaking Notes / Transcript",
    "quiz_exit_ticket": "Quiz / Exit Ticket",
    "homework_task": "Homework / Assignment",
    "general_work": "Classroom Work",
}

RETENTION_POLICIES = ("7_days", "30_days", "until_deleted", "manual_only")
RETENTION_LABELS = {
    "7_days": "7 Days (Short-term review)",
    "30_days": "30 Days (Recommended default)",
    "until_deleted": "Until Term End (Manual delete)",
    "manual_only": "Manual Only (Teacher controlled)",
}

SOURCE_FORMATS = ("pasted_text", "telegram_text", "txt_file", "docx_file")

DEFERRED_FORMATS = {
    ".pdf": "PDF text extraction is deferred until optical character and privacy safeguards are verified.",
    ".mp3": "Audio transcription is deferred until audio consent and privacy safeguards are verified.",
    ".m4a": "Audio transcription is deferred until audio consent and privacy safeguards are verified.",
    ".wav": "Audio transcription is deferred until audio consent and privacy safeguards are verified.",
    ".ogg": "Audio transcription is deferred until audio consent and privacy safeguards are verified.",
    ".jpg": "Image OCR is deferred until character recognition and consent verification are completed.",
    ".jpeg": "Image OCR is deferred until character recognition and consent verification are completed.",
    ".png": "Image OCR is deferred until character recognition and consent verification are completed.",
}

_STUDENT_HEADER = re.compile(
    r"(?im)^\s*(?:#+\s*)?(?:student|learner|pupil|response|sample|s|group|pair)\s*([a-z0-9_-]+)\s*[:.-]\s*"
)
_SEPARATOR_LINE = re.compile(r"(?im)^\s*[-=_*]{3,}\s*$")


class ClassEvidenceDisabledError(RuntimeError):
    pass


def _require_evidence() -> None:
    if not feature_enabled("evidence"):
        raise ClassEvidenceDisabledError("Evidence inbox is not enabled.")


def _utc_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.%fZ")


def _strip_control_chars(text: str) -> str:
    return "".join(c for c in text if c in ("\n", "\t") or ord(c) >= 32)


def sanitize_filename(filename: str | None) -> str | None:
    if not filename:
        return None
    basename = os.path.basename(str(filename)).strip()
    forbidden = set('<>:"/\\|?*')
    cleaned = "".join(c if (ord(c) >= 32 and c not in forbidden) else "_" for c in basename)
    cleaned = " ".join(cleaned.split()).strip()
    return cleaned[:120] if cleaned else "unnamed_file"


def parse_txt_bytes(data: bytes) -> str:
    if not data or len(data) == 0:
        raise ValueError("The uploaded text file is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise ValueError(f"File size exceeds the 2 MB limit ({len(data)} bytes).")
    if bytes([0]) in data:
        raise ValueError("Binary files cannot be processed as text evidence.")

    # CP1252 is the common Windows text encoding and must precede latin-1:
    # latin-1 accepts every byte and would otherwise decode smart quotes and
    # the euro sign incorrectly.
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            text = data.decode(encoding)
            cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
            if cleaned.strip():
                return cleaned
        except UnicodeDecodeError:
            continue
    raise ValueError("The text file could not be decoded. Ensure it is saved in UTF-8 format.")


def parse_docx_bytes(data: bytes) -> str:
    if not data or len(data) == 0:
        raise ValueError("The uploaded .docx file is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise ValueError(f"File size exceeds the 2 MB limit ({len(data)} bytes).")

    try:
        doc = Document(io.BytesIO(data))
    except (zipfile.BadZipFile, KeyError, Exception) as exc:
        raise ValueError(
            "The .docx file is corrupted, password-protected, or not a valid Word document."
        ) from exc

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    full_text = "\n\n".join(parts).strip()
    if not full_text:
        raise ValueError("The .docx document contains no readable text content.")
    return full_text


def split_evidence_text(raw_text: str) -> list[dict[str, str]]:
    if not isinstance(raw_text, str):
        raise ValueError("Evidence content must be text.")
    if len(raw_text.encode("utf-8")) > MAX_FILE_BYTES:
        raise ValueError(f"Pasted evidence exceeds the 2 MB limit ({len(raw_text.encode('utf-8'))} bytes).")
    text = " ".join(raw_text.splitlines(keepends=True))
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")

    items: list[dict[str, str]] = []
    current_label: str | None = None
    current_lines: list[str] = []

    def _flush(label: str | None, buffer: list[str], auto_idx: int) -> None:
        content = "\n".join(buffer).strip()
        content = _strip_control_chars(content)
        if len(content) >= MIN_ITEM_CHARS:
            final_label = (label or f"Student {auto_idx}").strip()
            items.append({
                "student_label": final_label[:80],
                "content": content[:MAX_ITEM_CHARS],
                "char_count": len(content[:MAX_ITEM_CHARS]),
                "word_count": len(content[:MAX_ITEM_CHARS].split()),
            })

    auto_counter = 1
    for line in lines:
        if _SEPARATOR_LINE.match(line):
            if current_lines:
                _flush(current_label, current_lines, auto_counter)
                auto_counter += 1
                current_lines = []
                current_label = None
            continue

        match = _STUDENT_HEADER.match(line)
        if match:
            # If current_label is None and current_lines has text, it's a preamble before first student
            if current_lines and current_label is not None:
                _flush(current_label, current_lines, auto_counter)
                auto_counter += 1
            current_lines = []
            tag = match.group(1).strip()
            current_label = f"Student {tag.upper() if tag.isalpha() and len(tag) <= 2 else tag}"
            line_body = line[match.end():].strip()
            if line_body:
                current_lines.append(line_body)
            continue

        current_lines.append(line)

    if current_lines:
        _flush(current_label, current_lines, auto_counter)

    # Fallback: if only 1 item and no explicit label, try double newlines
    if len(items) == 1 and not current_label:
        single_content = items[0]["content"]
        paragraphs = [p.strip() for p in single_content.split("\n\n") if len(p.strip()) >= 30]
        if len(paragraphs) > 1:
            items = []
            for idx, p in enumerate(paragraphs, 1):
                p_clean = _strip_control_chars(p)[:MAX_ITEM_CHARS]
                items.append({
                    "student_label": f"Student {idx}",
                    "content": p_clean,
                    "char_count": len(p_clean),
                    "word_count": len(p_clean.split()),
                })

    if not items:
        clean = _strip_control_chars(raw_text).strip()
        if len(clean) < MIN_ITEM_CHARS:
            raise ValueError("Evidence content is too short (minimum 2 characters required).")
        items.append({
            "student_label": "Student 1",
            "content": clean[:MAX_ITEM_CHARS],
            "char_count": len(clean[:MAX_ITEM_CHARS]),
            "word_count": len(clean[:MAX_ITEM_CHARS].split()),
        })

    if len(items) > MAX_BATCH_ITEMS:
        raise ValueError(
            f"Batch exceeds maximum limit of {MAX_BATCH_ITEMS} items (found {len(items)} items)."
        )

    return items


def validate_file_submission(
    filename: str, data: bytes
) -> tuple[str, str, list[dict[str, str]]]:
    _require_evidence()
    sanitized = sanitize_filename(filename) or "unnamed_file"
    ext = os.path.splitext(sanitized)[1].lower()

    if ext in DEFERRED_FORMATS:
        raise ValueError(DEFERRED_FORMATS[ext])

    if ext == ".txt":
        format_type = "txt_file"
        text = parse_txt_bytes(data)
    elif ext == ".docx":
        format_type = "docx_file"
        text = parse_docx_bytes(data)
    else:
        raise ValueError(
            f"Unsupported file format '{ext}'. TeacherOS accepts .txt and .docx files, or direct text paste."
        )

    items = split_evidence_text(text)
    return format_type, sanitized, items


def submit_evidence_batch(
    *,
    telegram_user: Any,
    class_id: int,
    evidence_type: str,
    raw_text: str | None = None,
    file_name: str | None = None,
    file_bytes: bytes | None = None,
    topic: str | None = None,
    lesson_id: int | None = None,
    objective_id: int | None = None,
    retention_policy: str = "30_days",
    privacy_confirmed: bool = True,
    database_path: Path | None = None,
) -> dict[str, Any]:
    """Safely ingest, validate, and persist an anonymous student evidence batch."""
    _require_evidence()
    if evidence_type not in EVIDENCE_TYPES:
        raise ValueError(f"Invalid evidence type '{evidence_type}'. Choose from {EVIDENCE_TYPES}.")
    if retention_policy not in RETENTION_POLICIES:
        raise ValueError(f"Invalid retention policy '{retention_policy}'.")
    if not privacy_confirmed:
        raise ValueError("Privacy confirmation is required before submitting student evidence.")

    if file_bytes is not None and file_name is not None:
        source_format, sanitized_filename, items = validate_file_submission(file_name, file_bytes)
    elif raw_text is not None:
        source_format = "pasted_text"
        sanitized_filename = None
        items = split_evidence_text(raw_text)
    else:
        raise ValueError("Provide either text content or an uploaded file.")

    if not items:
        raise ValueError("No valid student evidence items could be extracted.")

    batch_uuid = f"ev-batch-{secrets.token_hex(10)}"

    # Resolve plan limits before opening the write transaction. The entitlement
    # lookup initializes/opens SQLite itself; doing that while this connection
    # is active can deadlock SQLite in WAL mode.
    from entitlement_service import check_feature_access
    batch_access = check_feature_access(
        int(telegram_user.id),
        "evidence_batches_per_class",
        class_id=class_id,
        database_path=database_path,
    )
    if not batch_access["allowed"]:
        raise ValueError(batch_access.get("upgrade_prompt") or "Evidence batch limit reached.")
    item_limit_access = check_feature_access(
        int(telegram_user.id),
        "evidence_items_per_batch",
        class_id=class_id,
        database_path=database_path,
    )
    if item_limit_access.get("enforced") and item_limit_access.get("limit") is not None:
        if len(items) > int(item_limit_access["limit"]):
            raise ValueError(
                item_limit_access.get("upgrade_prompt")
                or f"This plan allows {item_limit_access['limit']} evidence items per batch."
            )

    with database_connection(database_path) as connection:
        user_id = ensure_database_user(connection, telegram_user)

        class_row = connection.execute(
            "SELECT id, display_name, status FROM classes WHERE id = ? AND user_id = ?",
            (class_id, user_id),
        ).fetchone()
        if class_row is None:
            raise ValueError("Class does not exist or does not belong to the user.")
        if class_row["status"] != "active":
            raise ValueError("Cannot submit evidence to an archived class.")

        if lesson_id is not None:
            lesson_check = connection.execute(
                "SELECT 1 FROM class_lessons WHERE id = ? AND class_id = ? AND user_id = ?",
                (lesson_id, class_id, user_id),
            ).fetchone()
            if lesson_check is None:
                raise ValueError("Linked lesson does not belong to this class.")

        if objective_id is not None:
            obj_check = connection.execute(
                "SELECT 1 FROM class_objectives WHERE id = ? AND class_id = ? AND user_id = ?",
                (objective_id, class_id, user_id),
            ).fetchone()
            if obj_check is None:
                raise ValueError("Linked objective does not belong to this class.")

        cursor = connection.execute(
            """
            INSERT INTO evidence_batches (
                batch_uuid, user_id, class_id, evidence_type, topic,
                lesson_id, objective_id, source_format, source_filename,
                item_count, retention_policy, privacy_confirmed, status
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'ready')
            """,
            (
                batch_uuid, user_id, class_id, evidence_type,
                topic.strip() if topic else None, lesson_id, objective_id,
                source_format, sanitized_filename, len(items),
                retention_policy, 1 if privacy_confirmed else 0,
            ),
        )
        batch_id = int(cursor.lastrowid)

        for item in items:
            connection.execute(
                """
                INSERT INTO evidence_items (
                    batch_id, class_id, user_id, student_label,
                    content, char_count, word_count, status
                ) VALUES (?, ?, ?, ?, ?, ?, ?, 'active')
                """,
                (
                    batch_id, class_id, user_id, item["student_label"],
                    item["content"], item["char_count"], item["word_count"],
                ),
            )

        # Emit privacy-compliant product event (ZERO raw text in telemetry)
        connection.execute(
            """
            INSERT OR IGNORE INTO product_events (
                event_uuid, user_id, class_id, event_name, privacy_class,
                properties_json, occurred_at
            ) VALUES (?, ?, ?, 'evidence_batch_created', 'product', ?, ?)
            """,
            (
                f"ev-event:{batch_uuid}",
                user_id,
                class_id,
                json.dumps({
                    "batch_id": batch_id,
                    "evidence_type": evidence_type,
                    "item_count": len(items),
                    "source_format": source_format,
                    "retention_policy": retention_policy,
                }, sort_keys=True),
                _utc_now(),
            ),
        )

    # Call get_evidence_batch AFTER exiting connection context
    return get_evidence_batch(
        telegram_user_id=int(telegram_user.id),
        batch_id=batch_id,
        database_path=database_path,
    )


def _resolve_telegram_user_id(
    telegram_user_id: int | None = None, telegram_user: Any = None
) -> int:
    if isinstance(telegram_user_id, int):
        return telegram_user_id
    if isinstance(telegram_user, int):
        return telegram_user
    if telegram_user is not None and hasattr(telegram_user, "id"):
        return int(telegram_user.id)
    raise ValueError("A valid telegram_user or telegram_user_id is required.")


def get_evidence_batch(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    batch_id: int,
    database_path: Path | None = None,
) -> dict[str, Any] | None:
    _require_evidence()
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)
    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            return None
        user_id = int(user["id"])

        batch = connection.execute(
            """
            SELECT b.*, c.display_name AS class_name, c.revision AS class_revision
            FROM evidence_batches AS b
            JOIN classes AS c ON c.id = b.class_id AND c.user_id = b.user_id
            WHERE b.id = ? AND b.user_id = ?
            """,
            (batch_id, user_id),
        ).fetchone()
        if batch is None:
            return None

        items = connection.execute(
            """
            SELECT id, batch_id, student_label, content, char_count, word_count,
                   status, created_at, updated_at
            FROM evidence_items
            WHERE batch_id = ? AND user_id = ? AND status = 'active'
            ORDER BY id ASC
            """,
            (batch_id, user_id),
        ).fetchall()

        batch_dict = dict(batch)
        batch_dict["items"] = [dict(row) for row in items]
        batch_dict["active_item_count"] = len(items)
        batch_dict["total_words"] = sum(int(item["word_count"]) for item in items)
        batch_dict["total_chars"] = sum(int(item["char_count"]) for item in items)
        return batch_dict


def list_evidence_batches(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    class_id: int,
    limit: int = 20,
    database_path: Path | None = None,
) -> list[dict[str, Any]]:
    _require_evidence()
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)
    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            return []
        user_id = int(user["id"])

        rows = connection.execute(
            """
            SELECT b.*, COUNT(i.id) AS active_items
            FROM evidence_batches AS b
            LEFT JOIN evidence_items AS i ON i.batch_id = b.id AND i.status = 'active'
            WHERE b.user_id = ? AND b.class_id = ? AND b.status != 'deleted'
            GROUP BY b.id
            ORDER BY b.created_at DESC, b.id DESC
            LIMIT ?
            """,
            (user_id, class_id, limit),
        ).fetchall()
        return [dict(r) for r in rows]


def update_evidence_item_label(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    item_id: int,
    new_label: str,
    database_path: Path | None = None,
) -> dict[str, Any] | None:
    _require_evidence()
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)
    cleaned = " ".join(str(new_label or "").split())
    if len(cleaned) < 1 or len(cleaned) > 80:
        raise ValueError("Student label must be between 1 and 80 characters.")

    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            return None
        user_id = int(user["id"])

        cursor = connection.execute(
            """
            UPDATE evidence_items
            SET student_label = ?, updated_at = ?
            WHERE id = ? AND user_id = ? AND status = 'active'
            """,
            (cleaned, _utc_now(), item_id, user_id),
        )
        if cursor.rowcount != 1:
            return None

        row = connection.execute(
            "SELECT * FROM evidence_items WHERE id = ? AND user_id = ?", (item_id, user_id)
        ).fetchone()
        return dict(row) if row is not None else None


def delete_evidence_item(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    item_id: int,
    database_path: Path | None = None,
) -> bool:
    """Soft-delete an individual evidence item with cascade count update."""
    _require_evidence()
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)
    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            return False
        user_id = int(user["id"])

        cursor = connection.execute(
            """
            UPDATE evidence_items
            SET status = 'deleted', deleted_at = ?, updated_at = ?
            WHERE id = ? AND user_id = ? AND status = 'active'
            """,
            (_utc_now(), _utc_now(), item_id, user_id),
        )
        return cursor.rowcount == 1


def delete_evidence_batch(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    batch_id: int,
    database_path: Path | None = None,
) -> bool:
    """Soft-delete an entire evidence batch and all of its items."""
    _require_evidence()
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)
    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            return False
        user_id = int(user["id"])

        now = _utc_now()
        cursor = connection.execute(
            """
            UPDATE evidence_batches
            SET status = 'deleted', deleted_at = ?, updated_at = ?
            WHERE id = ? AND user_id = ? AND status != 'deleted'
            """,
            (now, now, batch_id, user_id),
        )
        if cursor.rowcount != 1:
            return False

        connection.execute(
            """
            UPDATE evidence_items
            SET status = 'deleted', deleted_at = ?, updated_at = ?
            WHERE batch_id = ? AND user_id = ? AND status = 'active'
            """,
            (now, now, batch_id, user_id),
        )

        connection.execute(
            """
            INSERT OR IGNORE INTO product_events (
                event_uuid, user_id, event_name, privacy_class,
                properties_json, occurred_at
            ) VALUES (?, ?, 'evidence_batch_deleted', 'product', ?, ?)
            """,
            (
                f"ev-del:{batch_id}:{secrets.token_hex(6)}",
                user_id,
                json.dumps({"batch_id": batch_id}, sort_keys=True),
                now,
            ),
        )
        return True


def purge_expired_evidence(*, database_path: Path | None = None) -> int:
    """Permanently purge expired evidence batches and items based on retention policy."""
    with database_connection(database_path) as connection:
        now_utc = datetime.now(timezone.utc)
        purged_count = 0

        # 7-day retention
        seven_days_ago = (now_utc - timedelta(days=7)).strftime("%Y-%m-%dT%H:%M:%S.%fZ")
        # 30-day retention
        thirty_days_ago = (now_utc - timedelta(days=30)).strftime("%Y-%m-%dT%H:%M:%S.%fZ")

        cursor = connection.execute(
            """
            UPDATE evidence_items
            SET content = '[PURGED_BY_RETENTION_POLICY]', status = 'purged', updated_at = ?
            WHERE status != 'purged' AND batch_id IN (
                SELECT id FROM evidence_batches
                WHERE (retention_policy = '7_days' AND created_at <= ?)
                   OR (retention_policy = '30_days' AND created_at <= ?)
                   OR (status = 'deleted' AND deleted_at <= ?)
            )
            """,
            (_utc_now(), seven_days_ago, thirty_days_ago, seven_days_ago),
        )
        purged_count += cursor.rowcount

        connection.execute(
            """
            UPDATE evidence_batches
            SET status = 'purged', purged_at = ?, updated_at = ?
            WHERE status != 'purged' AND (
                (retention_policy = '7_days' AND created_at <= ?)
                OR (retention_policy = '30_days' AND created_at <= ?)
                OR (status = 'deleted' AND deleted_at <= ?)
            )
            """,
            (_utc_now(), _utc_now(), seven_days_ago, thirty_days_ago, seven_days_ago),
        )

        return purged_count
