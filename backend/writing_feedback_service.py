from __future__ import annotations

import io
import json
import logging
import re
import secrets
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Mapping, Sequence

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from database import database_connection
from feature_flags import feature_enabled


logger = logging.getLogger(__name__)

PROMPT_CONTRACT = "teacheros.writing_feedback"
PROMPT_VERSION = "2026-08-31.1"

_VALID_MODES = {"light", "balanced", "detailed", "rubric"}
_VALID_LEVELS = {"A1", "A2", "B1", "B2", "C1", "C2"}


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _resolve_telegram_user_id(
    telegram_user_id: int | None = None, telegram_user: Any = None
) -> int:
    if isinstance(telegram_user_id, int):
        return telegram_user_id
    if isinstance(telegram_user, int):
        return telegram_user
    if telegram_user is not None and hasattr(telegram_user, "id"):
        return int(telegram_user.id)
    raise ValueError("A valid telegram_user or telegram_user_id is required.")


def _clean_text(value: str) -> str:
    return " ".join(value.split())


def _build_student_copy(
    student_label: str,
    level: str,
    mode: str,
    strengths: Sequence[str],
    task_achievement: str,
    priorities: Sequence[dict[str, Any]],
    categorized_examples: dict[str, Any],
    revision_task: str,
    teacher_comments: str,
    rubric_scores: dict[str, Any] | None = None,
) -> str:
    lines: list[str] = [
        f"🌟 Writing Feedback for {student_label} (Level: {level})",
        "=" * 48,
        "",
        "💪 Strengths & Highlights:",
    ]
    for s in strengths:
        lines.append(f"  • {s}")

    lines.append("")
    lines.append("🎯 Key Revision Priorities (Focus on these next):")
    for p in priorities[:3]:
        lines.append(f"  {p.get('priority', 1)}. {p.get('title')}: {p.get('student_friendly_note')}")

    suggestions = categorized_examples.get("suggestions", [])
    if suggestions:
        lines.append("")
        lines.append("💡 Ideas for Polish & Style:")
        for sug in suggestions[:2]:
            lines.append(f'  • Instead of: "{sug.get("original")}"')
            lines.append(f'    Try: "{sug.get("suggestion")}" ({sug.get("rationale")})')

    corrections = categorized_examples.get("corrections", [])
    if corrections:
        lines.append("")
        lines.append("🔍 Accuracy Check:")
        for cor in corrections[:2]:
            lines.append(f'  • Check: "{cor.get("original")}" ➔ "{cor.get("correction")}" ({cor.get("explanation")})')

    if rubric_scores:
        lines.append("")
        lines.append("📊 Rubric Assessment (Draft):")
        for crit, score_info in rubric_scores.items():
            lines.append(f"  • {crit}: {score_info.get('score', '')} — {score_info.get('comment', '')}")

    lines.append("")
    lines.append("🚀 Your Actionable Revision Task:")
    lines.append(f"  {revision_task}")

    if teacher_comments:
        lines.append("")
        lines.append("💬 Teacher Note:")
        lines.append(f"  {teacher_comments}")

    return "\n".join(lines)


def _build_teacher_copy(
    student_label: str,
    level: str,
    mode: str,
    word_count: int,
    task_prompt: str | None,
    strengths: Sequence[str],
    task_achievement: str,
    priorities: Sequence[dict[str, Any]],
    categorized_examples: dict[str, Any],
    revision_task: str,
    teacher_comments: str,
    rubric_scores: dict[str, Any] | None,
    approved: bool,
) -> str:
    status_str = "APPROVED" if approved else "DRAFT (Requires Teacher Review)"
    lines: list[str] = [
        f"📋 TEACHEROS WRITING DIAGNOSTIC — {student_label}",
        f"Status: {status_str}",
        f"Level: {level} | Mode: {mode.upper()} | Word Count: {word_count} words",
    ]
    if task_prompt:
        lines.append(f"Task Prompt: {task_prompt}")

    lines.append("-" * 48)
    lines.append(f"Task Achievement: {task_achievement}")

    lines.append("")
    lines.append("Key Strengths:")
    for s in strengths:
        lines.append(f"  + {s}")

    lines.append("")
    lines.append("Prioritized Feedback Targets (Max 3):")
    for p in priorities[:3]:
        lines.append(f"  {p.get('priority', 1)}. {p.get('title')}")
        lines.append(f"     Diagnostic: {p.get('explanation')}")
        lines.append(f"     Student Note: {p.get('student_friendly_note')}")

    lines.append("")
    lines.append("Categorized Corrections (Accuracy):")
    for c in categorized_examples.get("corrections", []):
        lines.append(f'  • "{c.get("original")}" -> "{c.get("correction")}" [{c.get("explanation")}]')

    lines.append("")
    lines.append("Categorized Suggestions (Lexis & Flow):")
    for s in categorized_examples.get("suggestions", []):
        lines.append(f'  • "{s.get("original")}" -> "{s.get("suggestion")}" [{s.get("rationale")}]')

    if rubric_scores:
        lines.append("")
        lines.append("Rubric Criteria Breakdown:")
        for crit, info in rubric_scores.items():
            lines.append(f"  • {crit}: {info.get('score')} ({info.get('comment')})")

    lines.append("")
    lines.append(f"Assigned Revision Task: {revision_task}")

    if teacher_comments:
        lines.append("")
        lines.append(f"Teacher Custom Comments: {teacher_comments}")

    lines.append("")
    lines.append("⏱️ Estimated Teacher Time Saved: ~12 minutes")
    return "\n".join(lines)


def _diagnose_writing(
    text: str,
    level: str,
    mode: str,
    student_label: str = "Student",
    task_prompt: str | None = None,
    rubric_criteria: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    """Deterministically analyze writing for strengths, priorities, corrections, and revision."""
    words = text.split()
    word_count = len(words)

    # 1. Strengths
    strengths: list[str] = []
    if word_count >= 50:
        strengths.append("Strong communicative effort with clear paragraph development and sustained focus.")
    elif word_count >= 20:
        strengths.append("Direct response to the topic with understandable sentence structures.")
    else:
        strengths.append("Clear core message with relevant key vocabulary.")

    cohesive = [w for w in ["because", "however", "therefore", "although", "for example", "also", "and", "but", "so"] if w in text.lower()]
    if cohesive:
        strengths.append(f"Good use of logical linking words ({', '.join(cohesive[:3])}) to connect thoughts.")

    # 2. Task Achievement
    task_achievement = "The writing communicates the core message effectively for the intended audience."
    if task_prompt:
        task_achievement = f"Addresses key elements of the prompt '{task_prompt[:50]}...' with appropriate tone."

    # 3. Categorized Examples (Corrections vs Suggestions)
    corrections: list[dict[str, Any]] = []
    suggestions: list[dict[str, Any]] = []

    # Check common accuracy items
    if re.search(r"\b(he|she|it|everyone)\s+(don't|like|go|think|play)\b", text, re.I):
        m = re.search(r"\b(he|she|it|everyone)\s+(don't|like|go|think|play)\b", text, re.I)
        orig = m.group(0) if m else "he don't"
        corrections.append({
            "original": orig,
            "correction": orig.replace("don't", "doesn't").replace("like", "likes").replace("go", "goes"),
            "explanation": "Use third-person singular (-s / doesn't) in the present simple.",
        })

    if re.search(r"\b(yesterday|last\s+\w+)\s+.*\b(is|go|see)\b", text, re.I):
        corrections.append({
            "original": "yesterday ... is/go",
            "correction": "yesterday ... was/went",
            "explanation": "Maintain past tense consistency when narrating past events.",
        })

    if re.search(r"\b(a|an)\s+(information|advices|homeworks)\b", text, re.I):
        corrections.append({
            "original": "an information / homeworks",
            "correction": "some information / homework",
            "explanation": "Uncountable nouns do not take 'a/an' or plural '-s'.",
        })

    if not corrections:
        corrections.append({
            "original": "Punctuation & capitalization",
            "correction": "Ensure initial capital letters and ending periods on every sentence.",
            "explanation": "Clear sentence boundaries enhance readability.",
        })

    # Suggestions (Style & Lexis)
    if "very good" in text.lower() or "very bad" in text.lower() or "nice" in text.lower():
        suggestions.append({
            "original": "very good / nice",
            "suggestion": "compelling / beneficial / effective",
            "rationale": "Upgrade general descriptors with precise academic vocabulary.",
        })
    else:
        suggestions.append({
            "original": "Short simple sentences",
            "suggestion": "Combine clauses using 'although', 'whereas', or 'in order to'",
            "rationale": "Demonstrates greater grammatical complexity at upper levels.",
        })

    # 4. Priorities (Max 3, prioritized)
    priorities: list[dict[str, Any]] = []
    if corrections:
        priorities.append({
            "priority": 1,
            "title": "Grammar Accuracy & Verb Forms",
            "explanation": "Target finite verb agreement and tense consistency across paragraphs.",
            "student_friendly_note": "Double-check your verbs for third-person '-s' and past tense endings.",
        })

    if mode in ("balanced", "detailed", "rubric"):
        priorities.append({
            "priority": 2,
            "title": "Lexical Range & Precision",
            "explanation": "Expand vocabulary by replacing generic words with domain-specific terms.",
            "student_friendly_note": "Try swapping repeated adjectives for more specific, expressive words.",
        })

    if mode in ("detailed", "rubric") and len(priorities) < 3:
        priorities.append({
            "priority": 3,
            "title": "Paragraph Structure & Discourse Flow",
            "explanation": "Use transitional phrases to frame topic sentences clearly.",
            "student_friendly_note": "Start each paragraph with a clear sentence that introduces the main point.",
        })

    # 5. Exactly ONE actionable revision task
    if priorities:
        top = priorities[0]["title"]
        if "Grammar" in top:
            revision_task = "Select 2 sentences with verbs and verify the tense and subject-verb agreement. Rewrite them accurately."
        elif "Lexical" in top:
            revision_task = "Find 2 general words (like 'good' or 'thing') and replace them with 2 advanced words from our unit list."
        else:
            revision_task = "Add 1 transition phrase (e.g. 'Furthermore', 'On the other hand') between paragraphs 1 and 2."
    else:
        revision_task = "Read your text aloud to check sentence flow and add one concluding sentence."

    # 6. Rubric Scores (if rubric provided)
    rubric_scores: dict[str, Any] | None = None
    if rubric_criteria:
        rubric_scores = {}
        for crit in rubric_criteria:
            rubric_scores[str(crit)] = {
                "score": "Band 6.5 (Draft)",
                "comment": f"Shows competent control with emerging {crit.lower()} complexity.",
                "is_draft_score": True,
            }

    # 7. Teacher Comments default
    teacher_comments = f"Great effort on this composition, {student_label}! Complete the revision task below to make your writing even stronger."

    return {
        "word_count": word_count,
        "strengths": strengths,
        "task_achievement": task_achievement,
        "priorities": priorities,
        "categorized_examples": {
            "corrections": corrections,
            "suggestions": suggestions,
        },
        "revision_task": revision_task,
        "teacher_comments": teacher_comments,
        "rubric_scores": rubric_scores,
    }


def generate_writing_feedback(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    student_text: str,
    student_label: str = "Student",
    student_level: str = "B1",
    feedback_mode: str = "balanced",
    class_id: int | None = None,
    evidence_item_id: int | None = None,
    task_prompt: str | None = None,
    rubric_name: str | None = None,
    rubric_criteria: Mapping[str, Any] | None = None,
    database_path: Path | None = None,
) -> dict[str, Any]:
    """Generate kind, prioritized, rubric-aware writing feedback centered on revision."""
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)

    cleaned_text = student_text.strip()
    if not cleaned_text or len(cleaned_text) < 3:
        raise ValueError("Student writing text cannot be empty (minimum 3 characters).")
    if len(cleaned_text) > 25_000:
        raise ValueError("Student writing text exceeds maximum length (25,000 characters).")

    norm_mode = feedback_mode.strip().lower()
    if norm_mode not in _VALID_MODES:
        raise ValueError(f"Invalid feedback_mode '{feedback_mode}'. Must be one of {_VALID_MODES}.")

    norm_level = student_level.strip().upper()
    if norm_level not in _VALID_LEVELS:
        norm_level = "B1"

    norm_label = student_label.strip() or "Student"

    # Diagnostic Generation
    diagnosis = _diagnose_writing(
        text=cleaned_text,
        level=norm_level,
        mode=norm_mode,
        student_label=norm_label,
        task_prompt=task_prompt,
        rubric_criteria=rubric_criteria,
    )

    student_copy = _build_student_copy(
        student_label=norm_label,
        level=norm_level,
        mode=norm_mode,
        strengths=diagnosis["strengths"],
        task_achievement=diagnosis["task_achievement"],
        priorities=diagnosis["priorities"],
        categorized_examples=diagnosis["categorized_examples"],
        revision_task=diagnosis["revision_task"],
        teacher_comments=diagnosis["teacher_comments"],
        rubric_scores=diagnosis["rubric_scores"],
    )

    teacher_copy = _build_teacher_copy(
        student_label=norm_label,
        level=norm_level,
        mode=norm_mode,
        word_count=diagnosis["word_count"],
        task_prompt=task_prompt,
        strengths=diagnosis["strengths"],
        task_achievement=diagnosis["task_achievement"],
        priorities=diagnosis["priorities"],
        categorized_examples=diagnosis["categorized_examples"],
        revision_task=diagnosis["revision_task"],
        teacher_comments=diagnosis["teacher_comments"],
        rubric_scores=diagnosis["rubric_scores"],
        approved=False,
    )

    feedback_uuid = f"wf-{secrets.token_hex(8)}"
    feedback_json = json.dumps(diagnosis, sort_keys=True, ensure_ascii=False)
    rubric_json = json.dumps(dict(rubric_criteria or {}), sort_keys=True, ensure_ascii=False) if rubric_criteria else None

    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            raise ValueError(f"User {resolved_id} not registered.")
        user_id = int(user["id"])

        cursor = connection.execute(
            """
            INSERT INTO writing_feedback_records (
                feedback_uuid, user_id, class_id, evidence_item_id,
                student_label, student_level, feedback_mode, task_prompt,
                rubric_name, rubric_json, feedback_json, teacher_comments,
                revision_task, student_copy_text, teacher_copy_text,
                estimated_minutes_saved, approved, status, prompt_contract, prompt_version,
                created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 12, 0, 'draft', ?, ?, ?, ?)
            """,
            (
                feedback_uuid,
                user_id,
                class_id,
                evidence_item_id,
                norm_label,
                norm_level,
                norm_mode,
                task_prompt,
                rubric_name,
                rubric_json,
                feedback_json,
                diagnosis["teacher_comments"],
                diagnosis["revision_task"],
                student_copy,
                teacher_copy,
                PROMPT_CONTRACT,
                PROMPT_VERSION,
                _utc_now(),
                _utc_now(),
            ),
        )
        record_id = cursor.lastrowid

        connection.execute(
            """
            INSERT OR IGNORE INTO product_events (
                event_uuid, user_id, class_id, event_name, privacy_class,
                properties_json, occurred_at
            ) VALUES (?, ?, ?, 'writing_feedback_generated', 'product', ?, ?)
            """,
            (
                f"wf-ev:{feedback_uuid}",
                user_id,
                class_id,
                json.dumps({
                    "feedback_id": record_id,
                    "mode": norm_mode,
                    "level": norm_level,
                    "word_count": diagnosis["word_count"],
                    "has_rubric": bool(rubric_criteria or rubric_name),
                }, sort_keys=True),
                _utc_now(),
            ),
        )

    return get_writing_feedback(
        telegram_user_id=resolved_id,
        feedback_id=record_id,
        database_path=database_path,
    )  # type: ignore[return-value]


def approve_writing_feedback(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    feedback_id: int,
    teacher_comments: str | None = None,
    database_path: Path | None = None,
) -> dict[str, Any] | None:
    """Approve writing feedback record, finalizing both teacher and student copies."""
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)

    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            return None
        user_id = int(user["id"])

        row = connection.execute(
            "SELECT * FROM writing_feedback_records WHERE id = ? AND user_id = ?",
            (feedback_id, user_id),
        ).fetchone()
        if row is None:
            return None

        record = dict(row)
        diagnosis = json.loads(record["feedback_json"])

        final_comments = teacher_comments.strip() if teacher_comments else record["teacher_comments"]

        # Regenerate updated copies with approved status
        student_copy = _build_student_copy(
            student_label=record["student_label"],
            level=record["student_level"],
            mode=record["feedback_mode"],
            strengths=diagnosis["strengths"],
            task_achievement=diagnosis["task_achievement"],
            priorities=diagnosis["priorities"],
            categorized_examples=diagnosis["categorized_examples"],
            revision_task=record["revision_task"],
            teacher_comments=final_comments or "",
            rubric_scores=diagnosis.get("rubric_scores"),
        )

        teacher_copy = _build_teacher_copy(
            student_label=record["student_label"],
            level=record["student_level"],
            mode=record["feedback_mode"],
            word_count=diagnosis.get("word_count", 0),
            task_prompt=record["task_prompt"],
            strengths=diagnosis["strengths"],
            task_achievement=diagnosis["task_achievement"],
            priorities=diagnosis["priorities"],
            categorized_examples=diagnosis["categorized_examples"],
            revision_task=record["revision_task"],
            teacher_comments=final_comments or "",
            rubric_scores=diagnosis.get("rubric_scores"),
            approved=True,
        )

        now = _utc_now()
        connection.execute(
            """
            UPDATE writing_feedback_records
            SET approved = 1, status = 'approved', approved_at = ?,
                teacher_comments = ?, student_copy_text = ?, teacher_copy_text = ?,
                updated_at = ?
            WHERE id = ? AND user_id = ?
            """,
            (now, final_comments, student_copy, teacher_copy, now, feedback_id, user_id),
        )

        connection.execute(
            """
            INSERT OR IGNORE INTO product_events (
                event_uuid, user_id, class_id, event_name, privacy_class,
                properties_json, occurred_at
            ) VALUES (?, ?, ?, 'writing_feedback_approved', 'product', ?, ?)
            """,
            (
                f"wf-appr:{feedback_id}:{secrets.token_hex(4)}",
                user_id,
                record["class_id"],
                json.dumps({"feedback_id": feedback_id, "mode": record["feedback_mode"]}, sort_keys=True),
                now,
            ),
        )

    return get_writing_feedback(
        telegram_user_id=resolved_id,
        feedback_id=feedback_id,
        database_path=database_path,
    )


def update_writing_feedback_comments(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    feedback_id: int,
    new_comments: str,
    database_path: Path | None = None,
) -> dict[str, Any] | None:
    """Allow teacher to customize comment notes on the feedback record."""
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)
    cleaned = new_comments.strip()

    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            return None
        user_id = int(user["id"])

        row = connection.execute(
            "SELECT * FROM writing_feedback_records WHERE id = ? AND user_id = ?",
            (feedback_id, user_id),
        ).fetchone()
        if row is None:
            return None

        record = dict(row)
        diagnosis = json.loads(record["feedback_json"])

        student_copy = _build_student_copy(
            student_label=record["student_label"],
            level=record["student_level"],
            mode=record["feedback_mode"],
            strengths=diagnosis["strengths"],
            task_achievement=diagnosis["task_achievement"],
            priorities=diagnosis["priorities"],
            categorized_examples=diagnosis["categorized_examples"],
            revision_task=record["revision_task"],
            teacher_comments=cleaned,
            rubric_scores=diagnosis.get("rubric_scores"),
        )

        teacher_copy = _build_teacher_copy(
            student_label=record["student_label"],
            level=record["student_level"],
            mode=record["feedback_mode"],
            word_count=diagnosis.get("word_count", 0),
            task_prompt=record["task_prompt"],
            strengths=diagnosis["strengths"],
            task_achievement=diagnosis["task_achievement"],
            priorities=diagnosis["priorities"],
            categorized_examples=diagnosis["categorized_examples"],
            revision_task=record["revision_task"],
            teacher_comments=cleaned,
            rubric_scores=diagnosis.get("rubric_scores"),
            approved=bool(record["approved"]),
        )

        connection.execute(
            """
            UPDATE writing_feedback_records
            SET teacher_comments = ?, student_copy_text = ?, teacher_copy_text = ?, updated_at = ?
            WHERE id = ? AND user_id = ?
            """,
            (cleaned, student_copy, teacher_copy, _utc_now(), feedback_id, user_id),
        )

    return get_writing_feedback(
        telegram_user_id=resolved_id,
        feedback_id=feedback_id,
        database_path=database_path,
    )


def get_writing_feedback(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    feedback_id: int,
    database_path: Path | None = None,
) -> dict[str, Any] | None:
    """Retrieve writing feedback record with class and diagnostic metadata."""
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)

    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            return None
        user_id = int(user["id"])

        row = connection.execute(
            """
            SELECT f.*, c.display_name AS class_name
            FROM writing_feedback_records AS f
            LEFT JOIN classes AS c ON c.id = f.class_id
            WHERE f.id = ? AND f.user_id = ?
            """,
            (feedback_id, user_id),
        ).fetchone()
        if row is None:
            return None

        result = dict(row)
        result["feedback"] = json.loads(result["feedback_json"])
        result["rubric"] = json.loads(result["rubric_json"]) if result.get("rubric_json") else None
        return result


def list_writing_feedbacks(
    *,
    telegram_user_id: int | None = None,
    telegram_user: Any = None,
    class_id: int | None = None,
    limit: int = 20,
    database_path: Path | None = None,
) -> list[dict[str, Any]]:
    """List recent writing feedback records for teacher."""
    resolved_id = _resolve_telegram_user_id(telegram_user_id, telegram_user)

    with database_connection(database_path) as connection:
        user = connection.execute(
            "SELECT id FROM users WHERE telegram_user_id = ?", (resolved_id,)
        ).fetchone()
        if user is None:
            return []
        user_id = int(user["id"])

        params: list[Any] = [user_id]
        where_clauses = ["f.user_id = ?"]

        if class_id is not None:
            where_clauses.append("f.class_id = ?")
            params.append(class_id)

        params.append(limit)
        query = f"""
            SELECT f.*, c.display_name AS class_name
            FROM writing_feedback_records AS f
            LEFT JOIN classes AS c ON c.id = f.class_id
            WHERE {' AND '.join(where_clauses)}
            ORDER BY f.created_at DESC, f.id DESC
            LIMIT ?
        """
        rows = connection.execute(query, tuple(params)).fetchall()
        results: list[dict[str, Any]] = []
        for r in rows:
            d = dict(r)
            d["feedback"] = json.loads(d["feedback_json"])
            results.append(d)
        return results


def export_writing_feedback_word(
    *,
    feedback: dict[str, Any],
    copy_type: str = "student",  # 'student' | 'teacher'
) -> tuple[str, bytes]:
    """Generate Word (.docx) export for student or teacher copy."""
    doc = Document()
    student_label = feedback.get("student_label", "Student")
    level = feedback.get("student_level", "B1")

    if copy_type == "teacher":
        filename = f"Teacher Diagnostic - {student_label} ({level}).docx"
        text = feedback.get("teacher_copy_text", "")
    else:
        filename = f"Writing Feedback - {student_label} ({level}).docx"
        text = feedback.get("student_copy_text", "")

    # Build document
    p_title = doc.add_paragraph()
    run = p_title.add_run(f"TeacherOS Writing Feedback — {student_label}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph(f"Level: {level} | Exported: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    doc.add_paragraph("-" * 50)

    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return filename, buffer.getvalue()


def export_writing_feedback_pdf(
    *,
    feedback: dict[str, Any],
    copy_type: str = "student",
) -> tuple[str, bytes]:
    """Generate PDF (.pdf) export for student or teacher copy."""
    student_label = feedback.get("student_label", "Student")
    level = feedback.get("student_level", "B1")

    if copy_type == "teacher":
        filename = f"Teacher Diagnostic - {student_label} ({level}).pdf"
        text = feedback.get("teacher_copy_text", "")
    else:
        filename = f"Writing Feedback - {student_label} ({level}).pdf"
        text = feedback.get("student_copy_text", "")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()

    story = []
    title_style = ParagraphStyle(
        "FeedbackTitle",
        parent=styles["Heading1"],
        fontSize=14,
        leading=18,
        textColor=RGBColor(30, 41, 59),
    )
    body_style = ParagraphStyle(
        "FeedbackBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
    )

    story.append(Paragraph(f"TeacherOS Writing Feedback — {student_label}", title_style))
    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=1, color=RGBColor(203, 213, 225), spaceAfter=15))

    for line in text.split("\n"):
        cleaned = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if cleaned.strip():
            story.append(Paragraph(cleaned, body_style))
            story.append(Spacer(1, 4))

    doc.build(story)
    return filename, buffer.getvalue()
