"""TeacherOS Progress Report Service (Day 23).

Generates editable, evidence-safe class progress reports:
- Whole-class summaries, end-of-unit summaries, and teacher reflections.
- Strict evidence safety: Never invents attendance, effort, behavior, home support, or proficiency.
- Clear separation of Evidence, Teacher Judgment, and AI-supported wording.
- Edit/approval preview with versioned revision history.
- Share-safe Word and PDF export generation.
"""
from __future__ import annotations

import io
import json
import logging
import sqlite3
import uuid
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Sequence

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

from database import database_connection

logger = logging.getLogger(__name__)

VALID_REPORT_TYPES = frozenset({"whole_class_summary", "end_of_unit_summary", "teacher_reflection"})
VALID_EDITABLE_FIELDS = frozenset({
    "title",
    "learning_covered_text",
    "strengths_text",
    "priorities_text",
    "change_observed_text",
    "next_steps_text",
    "teacher_comments",
})

REPORT_TYPE_LABELS = {
    "whole_class_summary": "📊 Whole-Class Progress Summary",
    "end_of_unit_summary": "📖 End-of-Unit Learning Report",
    "teacher_reflection": "💡 Instructional Reflection & Review",
}


def _utc_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S.%fZ")


def _row_dict(row: sqlite3.Row | None) -> dict[str, Any] | None:
    if row is None:
        return None
    d = dict(row)
    if "evidence_summary_json" in d and isinstance(d["evidence_summary_json"], str):
        try:
            d["evidence_summary"] = json.loads(d["evidence_summary_json"])
        except Exception:
            d["evidence_summary"] = {}
    if "source_ids_json" in d and isinstance(d["source_ids_json"], str):
        try:
            d["source_ids"] = json.loads(d["source_ids_json"])
        except Exception:
            d["source_ids"] = []
    return d


# ---------------------------------------------------------------------------
# 1. Report Generation & Evidence Gathering
# ---------------------------------------------------------------------------

def generate_progress_report(
    *,
    user_id: int,
    class_id: int,
    report_type: str,
    reporting_period_start: str,
    reporting_period_end: str,
    unit_id: int | None = None,
    custom_title: str | None = None,
    database_path: Path | None = None,
) -> dict[str, Any] | None:
    """Generate an evidence-safe progress report grounded strictly in approved records."""
    if report_type not in VALID_REPORT_TYPES:
        raise ValueError(f"Invalid report type: {report_type}")
    try:
        period_start = date.fromisoformat(str(reporting_period_start).strip()[:10])
        period_end = date.fromisoformat(str(reporting_period_end).strip()[:10])
    except ValueError as exc:
        raise ValueError("Reporting period dates must use YYYY-MM-DD.") from exc
    if period_start > period_end:
        raise ValueError("Reporting period start must not be after the end.")
    reporting_period_start = period_start.isoformat()
    reporting_period_end = period_end.isoformat()

    now_str = _utc_now()
    report_uuid = f"rep_{uuid.uuid4().hex[:12]}"

    with database_connection(database_path) as conn:
        class_row = conn.execute(
            "SELECT * FROM classes WHERE id = ? AND user_id = ?",
            (class_id, user_id),
        ).fetchone()
        if not class_row:
            raise ValueError(f"Class {class_id} not found or user mismatch.")

        class_name = class_row["display_name"]
        level = class_row["level"]

        # Gather approved evidence records within period
        outcomes_rows = conn.execute(
            """
            SELECT o.*, l.title as lesson_title
            FROM lesson_outcomes AS o
            JOIN class_lessons AS l ON l.id = o.class_lesson_id
            WHERE o.class_id = ? AND o.user_id = ?
            AND o.status IN ('saved', 'approved')
            AND o.created_at BETWEEN ? AND ?
            ORDER BY o.created_at ASC
            """,
            (class_id, user_id, reporting_period_start, reporting_period_end + "T23:59:59.999Z"),
        ).fetchall()
        outcomes = [dict(r) for r in outcomes_rows]

        analyses_rows = conn.execute(
            """
            SELECT * FROM evidence_analysis_results
            WHERE class_id = ? AND user_id = ? AND approved = 1
            AND created_at BETWEEN ? AND ?
            ORDER BY created_at ASC
            """,
            (class_id, user_id, reporting_period_start, reporting_period_end + "T23:59:59.999Z"),
        ).fetchall()
        analyses = [dict(r) for r in analyses_rows]

        writing_feedbacks_rows = conn.execute(
            """
            SELECT * FROM writing_feedback_records
            WHERE class_id = ? AND user_id = ? AND approved = 1
            AND created_at BETWEEN ? AND ?
            ORDER BY created_at ASC
            """,
            (class_id, user_id, reporting_period_start, reporting_period_end + "T23:59:59.999Z"),
        ).fetchall()
        writing_feedbacks = [dict(r) for r in writing_feedbacks_rows]

        objectives_rows = conn.execute(
            """
            SELECT * FROM class_objectives
            WHERE class_id = ? AND user_id = ? AND status != 'archived'
            """,
            (class_id, user_id),
        ).fetchall()
        objectives = [dict(r) for r in objectives_rows]

        if unit_id is not None:
            unit_owner = conn.execute(
                "SELECT 1 FROM class_curriculum_units WHERE id = ? AND class_id = ? AND user_id = ?",
                (unit_id, class_id, user_id),
            ).fetchone()
            if unit_owner is None:
                raise ValueError("Curriculum unit does not belong to this class.")

        # Build source ID references
        sources: list[dict[str, Any]] = []
        for o in outcomes:
            sources.append({"type": "lesson_outcome", "id": o["id"]})
        for a in analyses:
            sources.append({"type": "evidence_analysis", "id": a["id"]})
        for w in writing_feedbacks:
            sources.append({"type": "writing_feedback", "id": w["id"]})
        for obj in objectives:
            sources.append({"type": "class_objective", "id": obj["id"]})

        total_evidence_count = len(outcomes) + len(analyses) + len(writing_feedbacks)
        has_insufficient_evidence = 1 if total_evidence_count == 0 else 0

        # Title formatting
        if custom_title and custom_title.strip():
            title = custom_title.strip()[:200]
        elif report_type == "end_of_unit_summary" and unit_id:
            unit = conn.execute("SELECT * FROM class_curriculum_units WHERE id = ?", (unit_id,)).fetchone()
            unit_name = f"Unit {unit['unit_number']}: {unit['unit_title']}" if unit else "Unit"
            title = f"{class_name} — {unit_name} Progress Report"
        elif report_type == "teacher_reflection":
            title = f"{class_name} — Instructional Reflection ({reporting_period_start} to {reporting_period_end})"
        else:
            title = f"{class_name} — Progress Summary ({reporting_period_start} to {reporting_period_end})"

        # Synthesize sections with honest boundary enforcement
        if has_insufficient_evidence:
            learning_covered_text = (
                "Insufficient recorded lesson evidence for this reporting period. "
                "Teacher observation or outcome check-in records are required."
            )
            strengths_text = "No approved student evidence records logged yet for this period."
            priorities_text = "Establish initial outcome checks and communicative evidence collection."
            change_observed_text = "Baseline period; no comparative change data recorded."
            next_steps_text = "Complete lesson outcome check-ins and log communicative evidence batches."
        else:
            # Learning covered
            covered_items = [f"• Lesson: {o.get('lesson_title') or 'Taught Lesson'} ({str(o.get('created_at', ''))[:10]})" for o in outcomes[:5]]
            secure_objs = [f"• Confirmed Target: {obj['objective']}" for obj in objectives if obj.get("is_secure")]
            learning_covered_text = "\n".join(covered_items + secure_objs) if (covered_items or secure_objs) else "Instruction aligned with active class syllabus."

            # Strengths
            strength_items = []
            for a in analyses[:3]:
                if a.get("approved_summary"):
                    strength_items.append(f"• Evidence finding: {a['approved_summary'][:120]}")
            for w in writing_feedbacks[:3]:
                # Keep report exports whole-class and privacy-safe; never
                # reproduce student labels even when a source record contains
                # an anonymized display token.
                strength_items.append("• Approved writing feedback record available for instructional planning")
            strengths_text = "\n".join(strength_items) if strength_items else "No specific strengths were recorded in the approved evidence for this period."

            # Priorities
            support_objs = [f"• Target needing scaffolding: {obj['objective']}" for obj in objectives if obj.get("status") == "needs_support"]
            priorities_text = "\n".join(support_objs) if support_objs else "Continue reinforcing target communicative can-do goals."

            # Change observed
            change_observed_text = f"Analyzed across {len(outcomes)} taught lessons and {len(analyses) + len(writing_feedbacks)} approved evidence items."

            # Next steps
            next_steps_text = "Prioritize guided review on identified support targets and conduct spaced retrieval."

        teacher_comments = "No teacher comments recorded."

        evidence_summary = {
            "outcomes_count": len(outcomes),
            "analyses_count": len(analyses),
            "writing_feedbacks_count": len(writing_feedbacks),
            "active_objectives_count": len(objectives),
            "has_insufficient_evidence": bool(has_insufficient_evidence),
        }

        cursor = conn.execute(
            """
            INSERT INTO class_progress_reports (
                report_uuid, user_id, class_id, report_type, title,
                reporting_period_start, reporting_period_end, unit_id,
                status, version, learning_covered_text, strengths_text,
                priorities_text, change_observed_text, next_steps_text,
                teacher_comments, has_insufficient_evidence,
                evidence_summary_json, source_ids_json, share_safe_verified,
                created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'draft', 1, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, ?, ?)
            """,
            (
                report_uuid, user_id, class_id, report_type, title,
                reporting_period_start, reporting_period_end, unit_id,
                learning_covered_text, strengths_text, priorities_text,
                change_observed_text, next_steps_text, teacher_comments,
                has_insufficient_evidence, json.dumps(evidence_summary),
                json.dumps(sources), now_str, now_str,
            ),
        )
        report_id = cursor.lastrowid
        row = conn.execute("SELECT * FROM class_progress_reports WHERE id = ?", (report_id,)).fetchone()
        return _row_dict(row)


# ---------------------------------------------------------------------------
# 2. Retrieval, Revision & Approval
# ---------------------------------------------------------------------------

def get_progress_report(
    *,
    user_id: int,
    report_id: int,
    database_path: Path | None = None,
) -> dict[str, Any] | None:
    """Retrieve single report with ownership verification."""
    with database_connection(database_path) as conn:
        row = conn.execute(
            """
            SELECT r.*, c.display_name AS class_name, c.level AS class_level
            FROM class_progress_reports AS r
            JOIN classes AS c ON c.id = r.class_id
            WHERE r.id = ? AND r.user_id = ?
            """,
            (report_id, user_id),
        ).fetchone()
        return _row_dict(row)


def list_progress_reports(
    *,
    user_id: int,
    class_id: int | None = None,
    status: str | None = None,
    limit: int = 20,
    database_path: Path | None = None,
) -> list[dict[str, Any]]:
    """List progress reports for a teacher / class."""
    with database_connection(database_path) as conn:
        where = ["r.user_id = ?"]
        params: list[Any] = [user_id]
        if class_id is not None:
            where.append("r.class_id = ?")
            params.append(class_id)
        if status is not None:
            where.append("r.status = ?")
            params.append(status)

        params.append(limit)
        query = f"""
            SELECT r.*, c.display_name AS class_name, c.level AS class_level
            FROM class_progress_reports AS r
            JOIN classes AS c ON c.id = r.class_id
            WHERE {' AND '.join(where)}
            ORDER BY r.updated_at DESC, r.id DESC
            LIMIT ?
        """
        rows = conn.execute(query, tuple(params)).fetchall()
        return [_row_dict(r) for r in rows if r is not None]


def update_progress_report_section(
    *,
    user_id: int,
    report_id: int,
    field_name: str,
    new_value: str,
    database_path: Path | None = None,
) -> dict[str, Any] | None:
    """Update a specific section of a report, increment version, and record revision history."""
    if field_name not in VALID_EDITABLE_FIELDS:
        raise ValueError(f"Invalid editable field: {field_name}")

    now_str = _utc_now()
    new_value = " ".join(str(new_value).split())

    with database_connection(database_path) as conn:
        report = conn.execute(
            "SELECT * FROM class_progress_reports WHERE id = ? AND user_id = ?",
            (report_id, user_id),
        ).fetchone()
        if not report:
            return None

        old_value = report[field_name]
        new_version = int(report["version"]) + 1

        # Record revision
        conn.execute(
            """
            INSERT INTO progress_report_revisions (
                report_id, user_id, version, field_changed, old_value, new_value, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (report_id, user_id, new_version, field_name, old_value, new_value, now_str),
        )

        # Any edit after approval invalidates the previous share-safe approval;
        # the teacher must explicitly approve the revised version again.
        approval_reset = ""
        if report["status"] == "approved":
            approval_reset = ", status = 'draft', share_safe_verified = 0, approved_at = NULL"

        # Update report
        conn.execute(
            f"""
            UPDATE class_progress_reports
            SET {field_name} = ?, version = ?, updated_at = ?{approval_reset}
            WHERE id = ? AND user_id = ?
            """,
            (new_value, new_version, now_str, report_id, user_id),
        )

        row = conn.execute("SELECT * FROM class_progress_reports WHERE id = ?", (report_id,)).fetchone()
        return _row_dict(row)


def approve_progress_report(
    *,
    user_id: int,
    report_id: int,
    database_path: Path | None = None,
) -> dict[str, Any] | None:
    """Explicit teacher approval gate marking report final and share-safe."""
    now_str = _utc_now()
    with database_connection(database_path) as conn:
        report = conn.execute(
            "SELECT * FROM class_progress_reports WHERE id = ? AND user_id = ?",
            (report_id, user_id),
        ).fetchone()
        if not report:
            return None

        conn.execute(
            """
            UPDATE class_progress_reports
            SET status = 'approved', share_safe_verified = 1, approved_at = ?, updated_at = ?
            WHERE id = ? AND user_id = ?
            """,
            (now_str, now_str, report_id, user_id),
        )

        row = conn.execute("SELECT * FROM class_progress_reports WHERE id = ?", (report_id,)).fetchone()
        return _row_dict(row)


def handle_deleted_source(
    *,
    source_type: str,
    source_id: int,
    database_path: Path | None = None,
) -> int:
    """Safe handling of deleted source data without corrupting existing approved reports."""
    with database_connection(database_path) as conn:
        reports = conn.execute("SELECT id, source_ids_json FROM class_progress_reports").fetchall()
        updated_count = 0
        now_str = _utc_now()
        for r in reports:
            try:
                sources = json.loads(r["source_ids_json"])
            except Exception:
                sources = []
            modified = False
            for s in sources:
                if s.get("type") == source_type and s.get("id") == source_id:
                    s["purged"] = True
                    modified = True
            if modified:
                conn.execute(
                    "UPDATE class_progress_reports SET source_ids_json = ?, updated_at = ? WHERE id = ?",
                    (json.dumps(sources), now_str, r["id"]),
                )
                updated_count += 1
        return updated_count


# ---------------------------------------------------------------------------
# 3. Export Word & PDF Generation
# ---------------------------------------------------------------------------

def _build_report_plain_text(report: dict[str, Any]) -> str:
    lines = [
        f"TeacherOS Progress Report — {report['title']}",
        "=" * 50,
        f"Report Type: {REPORT_TYPE_LABELS.get(report['report_type'], report['report_type'])}",
        f"Reporting Period: {report['reporting_period_start']} to {report['reporting_period_end']}",
        f"Status: {report['status'].upper()} (Version {report['version']})",
        "",
        "1. LEARNING COVERED & SYLLABUS TARGETS:",
        report.get("learning_covered_text", ""),
        "",
        "2. DEMONSTRATED STRENGTHS (Approved Evidence):",
        report.get("strengths_text", ""),
        "",
        "3. INSTRUCTIONAL PRIORITIES (Needing Support):",
        report.get("priorities_text", ""),
        "",
        "4. OBSERVED CHANGE & PROGRESS:",
        report.get("change_observed_text", ""),
        "",
        "5. NEXT STEPS & INSTRUCTIONAL RECOMMENDATIONS:",
        report.get("next_steps_text", ""),
        "",
        "6. TEACHER COMMENTS & PROFESSIONAL JUDGMENT:",
        report.get("teacher_comments", "None"),
        "",
        "🔒 Evidence Safety Notice: Grounded strictly in teacher-confirmed records and approved evidence analyses. No invented metrics.",
    ]
    return "\n".join(lines)


def export_progress_report_word(
    *,
    user_id: int,
    report_id: int,
    database_path: Path | None = None,
) -> tuple[str, bytes]:
    """Generate professional Word (.docx) export for an approved or preview report."""
    report = get_progress_report(user_id=user_id, report_id=report_id, database_path=database_path)
    if not report:
        raise ValueError("Report not found.")

    doc = Document()
    clean_title = report["title"].replace("/", "-").replace("\\", "-")
    filename = f"Progress Report - {clean_title} (v{report['version']}).docx"

    # Header
    p_title = doc.add_paragraph()
    run = p_title.add_run(report["title"])
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph(
        f"Type: {REPORT_TYPE_LABELS.get(report['report_type'], report['report_type'])} | "
        f"Period: {report['reporting_period_start']} to {report['reporting_period_end']} | "
        f"Status: {report['status'].upper()} (v{report['version']})"
    )
    doc.add_paragraph("-" * 60)

    sections = [
        ("Learning Covered", report.get("learning_covered_text", "")),
        ("Observed Strengths", report.get("strengths_text", "")),
        ("Instructional Priorities", report.get("priorities_text", "")),
        ("Observed Change", report.get("change_observed_text", "")),
        ("Next Instructional Steps", report.get("next_steps_text", "")),
        ("Teacher Comments", report.get("teacher_comments", "")),
    ]

    for sec_title, sec_content in sections:
        p_sec = doc.add_paragraph()
        r_sec = p_sec.add_run(sec_title)
        r_sec.font.size = Pt(12)
        r_sec.font.bold = True
        r_sec.font.color.rgb = RGBColor(51, 65, 85)

        for line in sec_content.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

    doc.add_paragraph()
    doc.add_paragraph("-" * 60)
    p_footer = doc.add_paragraph(
        "🔒 Evidence Safety: Generated by TeacherOS based strictly on teacher-confirmed records. "
        "No attendance, effort, or exact proficiency metrics were fabricated."
    )
    p_footer.runs[0].font.size = Pt(8)
    p_footer.runs[0].font.color.rgb = RGBColor(148, 163, 184)

    buffer = io.BytesIO()
    doc.save(buffer)
    return filename, buffer.getvalue()


def export_progress_report_pdf(
    *,
    user_id: int,
    report_id: int,
    database_path: Path | None = None,
) -> tuple[str, bytes]:
    """Generate professional PDF (.pdf) export for an approved or preview report."""
    report = get_progress_report(user_id=user_id, report_id=report_id, database_path=database_path)
    if not report:
        raise ValueError("Report not found.")

    clean_title = report["title"].replace("/", "-").replace("\\", "-")
    filename = f"Progress Report - {clean_title} (v{report['version']}).pdf"

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
    )
    styles = getSampleStyleSheet()

    story = []
    title_style = ParagraphStyle(
        "RepTitle",
        parent=styles["Heading1"],
        fontSize=15,
        leading=19,
        textColor=RGBColor(30, 41, 59),
    )
    h2_style = ParagraphStyle(
        "RepH2",
        parent=styles["Heading2"],
        fontSize=11,
        leading=15,
        textColor=RGBColor(51, 65, 85),
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "RepBody",
        parent=styles["Normal"],
        fontSize=9,
        leading=13,
        textColor=RGBColor(30, 41, 59),
    )

    story.append(Paragraph(f"TeacherOS Progress Report — {report['title']}", title_style))
    meta_text = (
        f"<b>Type:</b> {REPORT_TYPE_LABELS.get(report['report_type'], report['report_type'])} | "
        f"<b>Period:</b> {report['reporting_period_start']} to {report['reporting_period_end']} | "
        f"<b>Status:</b> {report['status'].upper()} (v{report['version']})"
    )
    story.append(Paragraph(meta_text, body_style))
    story.append(Spacer(1, 8))
    story.append(HRFlowable(width="100%", thickness=1, color=RGBColor(203, 213, 225), spaceAfter=10))

    sections = [
        ("1. Learning Covered & Syllabus Targets", report.get("learning_covered_text", "")),
        ("2. Observed Strengths (Approved Evidence)", report.get("strengths_text", "")),
        ("3. Instructional Priorities (Needing Support)", report.get("priorities_text", "")),
        ("4. Observed Change & Progress", report.get("change_observed_text", "")),
        ("5. Next Steps & Recommendations", report.get("next_steps_text", "")),
        ("6. Teacher Comments & Professional Judgment", report.get("teacher_comments", "")),
    ]

    for sec_title, sec_content in sections:
        story.append(Paragraph(sec_title, h2_style))
        for line in sec_content.split("\n"):
            cleaned = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if cleaned.strip():
                story.append(Paragraph(cleaned, body_style))
        story.append(Spacer(1, 6))

    story.append(HRFlowable(width="100%", thickness=1, color=RGBColor(203, 213, 225), spaceAfter=8))
    footer_text = (
        "🔒 <b>Evidence Safety:</b> Grounded strictly in teacher-confirmed records and approved evidence analyses. "
        "No attendance, effort, or exact proficiency metrics were fabricated."
    )
    story.append(Paragraph(footer_text, body_style))

    doc.build(story)
    return filename, buffer.getvalue()
